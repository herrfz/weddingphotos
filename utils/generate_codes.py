import sqlite3
import segno
import docx
from pathlib import Path
from docx.enum.table import WD_ALIGN_VERTICAL
from itertools import cycle


DBPATH = Path.cwd().parent / 'db.sqlite'
EVENT = 'oma'
BASEURL = 'https://weddingphotos-243848a36014.herokuapp.com'
DOCNAME = EVENT + '_tasks.docx'

tasks = None
users = None
doc = docx.Document()
table = doc.add_table(rows=0, cols=2)


try:
    with sqlite3.connect(DBPATH) as conn:
        cur = conn.cursor()
        tasks = cur.execute('SELECT id, description FROM tasks WHERE event = ?', (EVENT,)).fetchall()

        cur = conn.cursor()
        users = cur.execute('SELECT name FROM users').fetchall()
except sqlite3.Error as e:
    print(e)

if users is not None and tasks is not None:
    for user, task in zip(users, cycle(tasks)):
        taskid = str(task[0])
        taskurl = '/'.join([BASEURL, user[0], taskid])
        print(taskurl)

        qrcode = segno.make_qr(taskurl)
        filename = user[0] + '_' + taskid + '.png'
        qrcode.save(filename, scale=3, border=3)

        cells = table.add_row().cells
        # add image to first column
        par = cells[0].add_paragraph()
        run = par.add_run()
        run.add_picture(filename)
        # add description to second column
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].text = user[0]

    doc.save(DOCNAME)
